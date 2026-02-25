"""
Document upload and management routes.

Handles file upload, processing, vectorization, and document listing.
"""

import os
import uuid
import shutil
import logging
from pathlib import Path
from typing import List, Optional

from fastapi import APIRouter, Depends, UploadFile, File, HTTPException, BackgroundTasks
from pydantic import BaseModel

from backend.supabase_client import get_current_user  # noqa: F401 - used via Depends
from backend.db import (
    add_document, update_document_status, get_user_documents,
    get_all_documents, delete_document, get_document_stats,
)

router = APIRouter(prefix="/documents", tags=["documents"])
logger = logging.getLogger(__name__)

# Upload directory
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Supported file types
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".csv", ".json", ".py", ".js", ".ts", ".yaml", ".yml", ".html"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB


class DocumentResponse(BaseModel):
    id: int
    filename: str
    original_name: str
    file_type: str
    file_size: int
    knowledge_type: str
    chunk_count: int
    status: str
    uploaded_at: str


class DocumentStatsResponse(BaseModel):
    total_documents: int
    total_chunks: int
    by_type: dict


def process_and_index_document(doc_id: int, file_path: str, original_name: str, user_id: int):
    """
    Background task: Process uploaded file and add to Qdrant vector store.
    
    Steps:
    1. Load and parse the document
    2. Classify knowledge type
    3. Chunk the document
    4. Add chunks to Qdrant (with user_id)
    5. Update document status in DB
    """
    try:
        from langchain_core.documents import Document
        from ingestion.chunk_documents import DocumentChunker
        from vector_store.qdrant_store import QdrantVectorStore
        from knowledge.knowledge_classifier import KnowledgeClassifier
        from knowledge.decision_parser import DecisionParser

        # Step 1: Read file content
        ext = Path(file_path).suffix.lower()
        content = ""

        if ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                content = "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception:
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(file_path)
                docs = loader.load()
                content = "\n".join(d.page_content for d in docs)
        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(file_path)
                content = "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                with open(file_path, "r", errors="ignore") as f:
                    content = f.read()
        elif ext == ".csv":
            from langchain_community.document_loaders import CSVLoader
            loader = CSVLoader(file_path)
            docs = loader.load()
            content = "\n".join(d.page_content for d in docs)
        else:
            # Text-based files (txt, md, py, js, ts, yaml, html, json)
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

        if not content.strip():
            update_document_status(doc_id, "error")
            return

        # Step 2: Classify knowledge type
        classifier = KnowledgeClassifier()
        classification = classifier.classify(content, original_name)
        # Extract string from ClassificationResult
        if hasattr(classification, 'knowledge_type'):
            kt = classification.knowledge_type
            knowledge_type = str(kt.value) if hasattr(kt, 'value') else str(kt)
        else:
            knowledge_type = str(classification)

        # Step 3: Parse decision metadata if applicable
        metadata = {"source": original_name, "knowledge_type": knowledge_type}
        if knowledge_type == "decision":
            parser = DecisionParser()
            decision_meta = parser.parse(content)
            if hasattr(decision_meta, 'to_metadata'):
                raw = decision_meta.to_metadata()
            elif isinstance(decision_meta, dict):
                raw = decision_meta
            else:
                raw = {}
            # Sanitize metadata values
            for k, v in raw.items():
                if v is None:
                    continue
                elif isinstance(v, (list, tuple)):
                    metadata[k] = ", ".join(str(i) for i in v) if v else ""
                elif isinstance(v, bool):
                    metadata[k] = str(v)
                else:
                    metadata[k] = str(v)

        # Step 4: Create LangChain document and chunk
        doc = Document(page_content=content, metadata=metadata)
        chunker = DocumentChunker()
        chunks = chunker.chunk([doc])

        # Step 5: Add to Qdrant vector store (per-user)
        qdrant = QdrantVectorStore()
        qdrant.add_documents(chunks, user_id=user_id)

        # Step 6: Update DB
        update_document_status(doc_id, "indexed", len(chunks), knowledge_type)

        logger.info(f"Document '{original_name}' indexed: {len(chunks)} chunks for user {user_id}")

    except Exception as e:
        print(f"[ERROR] Document processing failed for {original_name}: {e}")
        import traceback
        traceback.print_exc()
        update_document_status(doc_id, "error")


@router.post("/upload", response_model=List[DocumentResponse])
async def upload_documents(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    user=Depends(get_current_user),
):
    """
    Upload one or more documents for knowledge base indexing.
    
    Supports: PDF, TXT, MD, DOCX, CSV, JSON, PY, JS, TS, YAML, HTML
    Max size: 20MB per file
    """
    results = []

    for file in files:
        # Validate extension
        ext = Path(file.filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
            )

        # Read file content
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail=f"File {file.filename} exceeds 20MB limit")

        # Save file
        unique_name = f"{uuid.uuid4().hex}_{file.filename}"
        file_path = os.path.join(UPLOAD_DIR, unique_name)
        with open(file_path, "wb") as f:
            f.write(content)

        # Record in database
        doc_id = add_document(
            user_id=user["id"],
            filename=unique_name,
            original_name=file.filename,
            file_type=ext.lstrip("."),
            file_size=len(content),
        )

        # Schedule background processing (pass user_id for Qdrant per-user isolation)
        background_tasks.add_task(process_and_index_document, doc_id, file_path, file.filename, user["id"])

        results.append(DocumentResponse(
            id=doc_id,
            filename=unique_name,
            original_name=file.filename,
            file_type=ext.lstrip("."),
            file_size=len(content),
            knowledge_type="pending",
            chunk_count=0,
            status="processing",
            uploaded_at="just now",
        ))

    return results


@router.get("/", response_model=List[DocumentResponse])
async def list_documents(user=Depends(get_current_user)):
    """List documents uploaded by the current user."""
    docs = get_user_documents(user["id"])
    return [
        DocumentResponse(
            id=d["id"],
            filename=d["filename"],
            original_name=d["original_name"],
            file_type=d["file_type"],
            file_size=d["file_size"],
            knowledge_type=d["knowledge_type"] or "explicit",
            chunk_count=d["chunk_count"] or 0,
            status=d["status"],
            uploaded_at=d["uploaded_at"] or "",
        )
        for d in docs
    ]


@router.get("/my", response_model=List[DocumentResponse])
async def list_my_documents(user=Depends(get_current_user)):
    """List documents uploaded by the current user."""
    docs = get_user_documents(user["id"])
    return [
        DocumentResponse(
            id=d["id"],
            filename=d["filename"],
            original_name=d["original_name"],
            file_type=d["file_type"],
            file_size=d["file_size"],
            knowledge_type=d["knowledge_type"] or "explicit",
            chunk_count=d["chunk_count"] or 0,
            status=d["status"],
            uploaded_at=d["uploaded_at"] or "",
        )
        for d in docs
    ]


@router.delete("/{doc_id}")
async def remove_document(doc_id: int, user=Depends(get_current_user)):
    """Delete a document (only your own uploads) and its vectors."""
    # Get the document to know its original_name for vector deletion
    docs = get_user_documents(user["id"])
    doc_record = next((d for d in docs if d["id"] == doc_id), None)

    success = delete_document(doc_id, user["id"])
    if not success:
        raise HTTPException(status_code=404, detail="Document not found or not owned by you")

    # Remove vectors from Qdrant
    if doc_record:
        try:
            from vector_store.qdrant_store import QdrantVectorStore
            qdrant = QdrantVectorStore()
            qdrant.delete_by_source(user["id"], doc_record["original_name"])
        except Exception as e:
            logger.warning(f"Failed to delete vectors for doc {doc_id}: {e}")

    return {"status": "deleted", "id": doc_id}


@router.get("/stats", response_model=DocumentStatsResponse)
async def document_stats():
    """Get document statistics for the knowledge base."""
    return get_document_stats()
